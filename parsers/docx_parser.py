from docx import Document

def parse_docx(file_path):
    try:
        doc = Document(file_path)
        text_parts = []

        # Обычные параграфы
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

        # Текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)

        full_text = "\n".join(text_parts)
        print(f"[parse_docx] Извлечено символов: {len(full_text)}")
        return full_text

    except Exception as e:
        print(f"[parse_docx] Ошибка при разборе DOCX: {e}")
        return ""
